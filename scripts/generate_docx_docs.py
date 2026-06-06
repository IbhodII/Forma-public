#!/usr/bin/env python3

# -*- coding: utf-8 -*-

"""

Генерация документации Word (.docx) из Markdown в docs/.



Запуск из корня проекта:

    python scripts/generate_docx_docs.py



Требуется: pip install python-docx

"""

from __future__ import annotations



import re

import subprocess

import sys

from pathlib import Path



ROOT = Path(__file__).resolve().parent.parent

DOCS = ROOT / "docs"

OUT = DOCS / "docx"





def ensure_python_docx() -> None:

    try:

        import docx  # noqa: F401

    except ImportError:

        print("Installing python-docx...")

        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])





def md_to_docx(md_path: Path, docx_path: Path, *, title: str | None = None) -> None:

    from docx import Document

    from docx.shared import Pt

    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



    text = md_path.read_text(encoding="utf-8")

    doc = Document()

    style = doc.styles["Normal"]

    style.font.name = "Calibri"

    style.font.size = Pt(11)



    if title:

        h = doc.add_heading(title, level=0)

        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT



    in_code = False

    in_mermaid = False

    code_lines: list[str] = []

    table_rows: list[list[str]] = []



    def flush_table() -> None:

        nonlocal table_rows

        if not table_rows:

            return

        cols = max(len(r) for r in table_rows)

        tbl = doc.add_table(rows=len(table_rows), cols=cols)

        tbl.style = "Table Grid"

        for ri, row in enumerate(table_rows):

            for ci, cell in enumerate(row):

                if ci < cols:

                    tbl.rows[ri].cells[ci].text = cell.strip()

        table_rows = []

        doc.add_paragraph()



    lines = text.splitlines()

    i = 0

    while i < len(lines):

        line = lines[i]



        if line.strip().startswith("```"):

            fence = line.strip()

            if not in_code:

                in_code = True

                in_mermaid = "mermaid" in fence

                code_lines = []

            else:

                if not in_mermaid and code_lines:

                    p = doc.add_paragraph()

                    run = p.add_run("\n".join(code_lines))

                    run.font.name = "Consolas"

                    run.font.size = Pt(9)

                in_code = False

                in_mermaid = False

                code_lines = []

            i += 1

            continue



        if in_code:

            if not in_mermaid:

                code_lines.append(line)

            i += 1

            continue



        if line.strip().startswith("|") and "|" in line[1:]:

            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):

                i += 1

                continue

            cells = [c.strip() for c in line.strip().strip("|").split("|")]

            table_rows.append(cells)

            i += 1

            continue

        flush_table()



        if line.strip() == "---":

            doc.add_paragraph("—" * 40)

            i += 1

            continue



        m = re.match(r"^(#{1,4})\s+(.*)$", line)

        if m:

            level = min(len(m.group(1)), 4)

            doc.add_heading(m.group(2).strip(), level=level)

            i += 1

            continue



        if line.strip().startswith("- "):

            doc.add_paragraph(line.strip()[2:], style="List Bullet")

            i += 1

            continue



        if line.strip().startswith(">"):

            p = doc.add_paragraph(line.strip().lstrip("> ").strip())

            p.paragraph_format.left_indent = Pt(18)

            i += 1

            continue



        if line.strip():

            doc.add_paragraph(line.strip())

        i += 1



    flush_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(docx_path))

    print(f"  OK  {docx_path.name}")





def combine_md(paths: list[Path], header: str) -> Path:

    """Временный объединённый .md для одного docx."""

    tmp = OUT / "_combined_temp.md"

    parts = [f"# {header}\n"]

    for p in paths:

        parts.append(f"\n\n---\n\n")

        parts.append(p.read_text(encoding="utf-8"))

    tmp.write_text("".join(parts), encoding="utf-8")

    return tmp





def main() -> None:

    ensure_python_docx()

    OUT.mkdir(parents=True, exist_ok=True)

    print(f"Generating .docx -> {OUT}\n")



    jobs: list[tuple[Path, Path, str | None]] = [

        (DOCS / "DESKTOP_IMPROVEMENTS.md", OUT / "MyHealthDashboard_Десктоп.docx", "Десктоп Forma (Electron)"),

        (DOCS / "PROJECT_CONTEXT.md", OUT / "MyHealthDashboard_Обзор_проекта.docx", None),

        (DOCS / "WORKOUT_PRESETS.md", OUT / "MyHealthDashboard_Пресеты_тренировок.docx", None),

        (DOCS / "STRETCHING.md", OUT / "MyHealthDashboard_Растяжка.docx", "Растяжка"),

        (DOCS / "BIKE.md", OUT / "MyHealthDashboard_Велотренировки.docx", "Велотренировки и FIT"),

        (DOCS / "NUTRITION.md", OUT / "MyHealthDashboard_Питание.docx", "Питание"),
        (
            DOCS / "UNITS_CONVERSION.md",
            OUT / "MyHealthDashboard_Единицы_измерения.docx",
            "Единицы измерения (метрика → american)",
        ),

    ]



    combined = combine_md(

        [
            DOCS / "API.md",
            DOCS / "DATABASE.md",
            DOCS / "SERVICES.md",
            DOCS / "CHANGELOG.md",
        ],

        "API, база данных и сервисы",

    )

    jobs.append((combined, OUT / "MyHealthDashboard_API_и_данные.docx", None))



    for src, dst, title in jobs:

        if not src.is_file():

            print(f"  SKIP (нет файла) {src.name}")

            continue

        md_to_docx(src, dst, title=title)



    tmp = OUT / "_combined_temp.md"

    if tmp.is_file():

        tmp.unlink()



    print("\nDone.")





if __name__ == "__main__":

    main()

